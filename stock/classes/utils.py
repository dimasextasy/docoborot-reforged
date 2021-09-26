import os
import zipfile

from Crypto.Hash import SHA256
from Crypto.PublicKey import RSA
from Crypto.Signature import pkcs1_15
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def generate_report(data):
    document = Document()

    head = document.add_heading(
        'Квитанция на оплату закупок компанией "{}" с {} по {}'.format(data['Компания'], data['Дата начала отсчета'],
                                                                       data['Дата конца отсчета']), 0)
    head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p = document.add_paragraph('Уважаемый(ая) ')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE
    p.add_run(data['Фамилия'] + ' ' + data['Имя'] + ' ' + data['Отчество'] + ' ').bold = True
    p.add_run('просим вас оплатить закупки, произведенные компанией ')
    p.add_run('"' + data['Компания'] + '"' + ', ').bold = True
    p.add_run('в которой для нас вы являеетесь контактным лицом в период с ')
    p.add_run(str(data['Дата начала отсчета']) + ' ').bold = True
    p.add_run('по ')
    p.add_run(str(data['Дата конца отсчета']) + '.').bold = True

    document.add_paragraph()
    document.add_paragraph('Прилагаю список товаров:')

    records = data['Информация о закупках']

    table = document.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Дата'
    hdr_cells[1].text = 'Продукт'
    hdr_cells[2].text = 'Количество'
    hdr_cells[3].text = 'Цена'
    hdr_cells[4].text = 'Итого'

    for deal in records.keys():
        print('это новое {}'.format(deal))
        row_cells = table.add_row().cells
        row_cells[0].text = str(deal)
        row_cells[1].text = records[deal][0]['Продукт']
        row_cells[2].text = records[deal][0]['Количество'].__str__()
        row_cells[3].text = records[deal][0]['Цена продажи'].__str__()
        row_cells[4].text = records[deal][0]['Сумма'].__str__()
        for index, element in enumerate(records[deal]):
            if index > 0:
                row_cells = table.add_row().cells
                row_cells[1].text = records[deal][index]['Продукт']
                row_cells[2].text = records[deal][index]['Количество'].__str__()
                row_cells[3].text = records[deal][index]['Цена продажи'].__str__()
                row_cells[4].text = records[deal][index]['Сумма'].__str__()

    document.add_paragraph()
    p = document.add_paragraph('Итого общая сумма вашей задолженности составляет ')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE
    p.add_run(data['Общая сумма закупок'].__str__() + ' рублей').bold = True

    document.save('Отчет.docx')

    create_signature()
    archive_files()


def create_signature():
    # Генерируете новый ключ (или берете ранее сгенерированный)
    key = RSA.generate(1024, os.urandom)
    # Получаете хэш файла
    file_hash = SHA256.new()
    with open("Отчет.docx", "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            file_hash.update(chunk)

    # Подписываете хэш
    signature = pkcs1_15.new(key).sign(file_hash)

    # Получаете открытый ключ из закрытого
    pubkey = key.publickey()

    pkcs1_15.new(pubkey).verify(file_hash, signature)

    with open('Подпись.txt', 'wb') as f:
        f.write(signature)
        f.close()

    with open('Ключ.pem', 'wb') as f:
        f.write(pubkey.export_key())
        f.close()


def archive_files():
    zip_file = zipfile.ZipFile('media/report.zip', 'w')
    files = ['Отчет.docx', 'Подпись.txt', 'Ключ.pem']
    for file in files:
        zip_file.write(file, compress_type=zipfile.ZIP_DEFLATED)
    zip_file.close()

    for file in files:
        os.remove(file)

